# import libraries
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml.ns import qn
import pickle
import win32com
from self import self
from win32com.client import Dispatch
import tempfile


import win32api
import win32print
import os, sys

document = Document()

# set up page to landscape
section = document.sections[0]
section.orientation = WD_ORIENTATION.LANDSCAPE

# set the page size to A4
section.page_width = Cm(29.7)
section.page_height = Cm(21)

# set up page margins
section.left_margin, section.right_margin = Inches(0.4), Inches(0.4)
section.top_margin, section.bottom_margin = Inches(0.2), Inches(0.1)

jobNumber = 4000

# Add a new table
table = document.add_table(rows=13, cols=2, style='Table Grid')

# Change row height
table.rows[0].height = Cm(1)
table.rows[1].height = Cm(1)
table.rows[2].height = Cm(1)
table.rows[7].height = Cm(1)
table.rows[8].height = Cm(1)
table.rows[10].height = Cm(1)

# Merge cells for customer requirement
table.cell(2, 0).merge(table.cell(3, 0))
table.cell(4, 0).merge(table.cell(6, 0))
table.cell(8, 0).merge(table.cell(9, 0))
table.cell(10, 0).merge(table.cell(12, 0))
table.cell(0, 0).merge(table.cell(0, 1))

table.cell(1, 1).merge(table.cell(3, 1))
table.cell(4, 1).merge(table.cell(6, 1))
table.cell(7, 1).merge(table.cell(9, 1))
table.cell(11, 1).merge(table.cell(12, 1))

# Add text to each cells
table.cell(0, 0).text = "Job Number " + str(jobNumber)
table.cell(1, 0).text = "Date:"
table.cell(3, 0).text = "Client:"
table.cell(5, 0).text = "Phone & Email:"
table.cell(7, 0).text = "Password:"
table.cell(9, 0).text = "Parts:"
table.cell(10, 0).text = "Items Serviced:"

table.cell(1, 1).text = "Address:"
table.cell(4, 1).text = "Issue:"
table.cell(7, 1).text = "Work Done:" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "Data Saved? Y / N"
table.cell(10, 1).text = "Misc Notes:"
table.cell(11, 1).text = "To Invoice:"

# Add a new column, split in the middle
section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')

# Add a space between two tables
paragraph = document.add_paragraph()

jobNumber +=1

table = document.add_table(rows=13, cols=2, style='Table Grid')
table.rows[0].height = Cm(1)
table.rows[1].height = Cm(1)
table.rows[2].height = Cm(1)
table.rows[7].height = Cm(1)
table.rows[8].height = Cm(1)
table.rows[10].height = Cm(1)

table.cell(2, 0).merge(table.cell(3, 0))
table.cell(4, 0).merge(table.cell(6, 0))
table.cell(8, 0).merge(table.cell(9, 0))
table.cell(10, 0).merge(table.cell(12, 0))
table.cell(0, 0).merge(table.cell(0, 1))

table.cell(0, 0).text = "Job Number: " + str(jobNumber)
table.cell(1, 0).text = "Date:"
table.cell(3, 0).text = "Client:"
table.cell(5, 0).text = "Phone & Email:"
table.cell(7, 0).text = "Password:"
table.cell(9, 0).text = "Parts:"
table.cell(10, 0).text = "Items Serviced:"

table.cell(1, 1).merge(table.cell(3, 1))
table.cell(4, 1).merge(table.cell(6, 1))
table.cell(7, 1).merge(table.cell(9, 1))
table.cell(11, 1).merge(table.cell(12, 1))

table.cell(1, 1).text = "Address:"
table.cell(4, 1).text = "Issue:"
table.cell(7, 1).text = "Work Done:" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "Data Saved? Y / N"
table.cell(10, 1).text = "Misc Notes:"
table.cell(11, 1).text = "To Invoice:"

# paragraph = document.add_paragraph()

jobNumber +=1
table = document.add_table(rows=13, cols=2, style='Table Grid')
table.rows[0].height = Cm(1)
table.rows[1].height = Cm(1)
table.rows[2].height = Cm(1)
table.rows[7].height = Cm(1)
table.rows[8].height = Cm(1)
table.rows[10].height = Cm(1)

table.cell(2, 0).merge(table.cell(3, 0))
table.cell(4, 0).merge(table.cell(6, 0))
table.cell(8, 0).merge(table.cell(9, 0))
table.cell(10, 0).merge(table.cell(12, 0))
table.cell(0, 0).merge(table.cell(0, 1))

table.cell(0, 0).text = "Job Number: " + str(jobNumber)
table.cell(1, 0).text = "Date:"
table.cell(3, 0).text = "Client:"
table.cell(5, 0).text = "Phone & Email:"
table.cell(7, 0).text = "Password:"
table.cell(9, 0).text = "Parts:"
table.cell(10, 0).text = "Items Serviced:"

table.cell(1, 1).merge(table.cell(3, 1))
table.cell(4, 1).merge(table.cell(6, 1))
table.cell(7, 1).merge(table.cell(9, 1))
table.cell(11, 1).merge(table.cell(12, 1))

table.cell(1, 1).text = "Address:"
table.cell(4, 1).text = "Issue:"
table.cell(7, 1).text = "Work Done:" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "Data Saved? Y / N"
table.cell(10, 1).text = "Misc Notes:"
table.cell(11, 1).text = "To Invoice:"

paragraph = document.add_paragraph()

jobNumber +=1
table = document.add_table(rows=13, cols=2, style='Table Grid')
table.rows[0].height = Cm(1)
table.rows[1].height = Cm(1)
table.rows[2].height = Cm(1)
table.rows[7].height = Cm(1)
table.rows[8].height = Cm(1)
table.rows[10].height = Cm(1)

table.cell(2, 0).merge(table.cell(3, 0))
table.cell(4, 0).merge(table.cell(6, 0))
table.cell(8, 0).merge(table.cell(9, 0))
table.cell(10, 0).merge(table.cell(12, 0))
table.cell(0, 0).merge(table.cell(0, 1))

table.cell(0, 0).text = "Job Number: " + str(jobNumber)
table.cell(1, 0).text = "Date:"
table.cell(3, 0).text = "Client:"
table.cell(5, 0).text = "Phone & Email:"
table.cell(7, 0).text = "Password:"
table.cell(9, 0).text = "Parts:"
table.cell(10, 0).text = "Items Serviced:"

table.cell(1, 1).merge(table.cell(3, 1))
table.cell(4, 1).merge(table.cell(6, 1))
table.cell(7, 1).merge(table.cell(9, 1))
table.cell(11, 1).merge(table.cell(12, 1))

table.cell(1, 1).text = "Address:"
table.cell(4, 1).text = "Issue:"
table.cell(7, 1).text = "Work Done:" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "Data Saved? Y / N"
table.cell(10, 1).text = "Misc Notes:"
table.cell(11, 1).text = "To Invoice:"

# save jobNumber in file.
file = open('var', 'wb')
pickle.dump(jobNumber, file)
file.close()

document.save('JobCard.docx')


path = ("JobCard.docx")
win32api.ShellExecute (
 0,
 "printto",
 path,
 '"%s"' % win32print.GetDefaultPrinter (),
 ".",
 0
)

