# import libraries
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Cm
from docx.oxml.ns import qn
import pickle

document = Document()

# set up page to landscape
section = document.sections[0]
section.orientation = WD_ORIENTATION.LANDSCAPE

# set the page size to A4
section.page_width = Cm(29.7)
section.page_height = Cm(21)

# set up page margins
section.left_margin, section.right_margin = Cm(1), Cm(1)
section.top_margin, section.bottom_margin = Cm(0.7), Cm(0.7)

jobNumber = 4060


# Add a new table
table = document.add_table(rows=8, cols=2, style='Table Grid')

# Change row height
table.rows[0].height = Cm(1.1)
table.rows[1].height = Cm(1.1)
table.rows[2].height = Cm(1.1)
table.rows[3].height = Cm(1.1)
table.rows[4].height = Cm(1.1)
table.rows[5].height = Cm(1.1)
table.rows[6].height = Cm(1.1)
table.rows[7].height = Cm(1.1)

table.columns[0].width = Cm(6.8)
table.columns[1].width = Cm(6.8)

table.cell(1, 0).merge(table.cell(2, 0))
table.cell(4, 0).merge(table.cell(4, 1))
table.cell(5, 1).merge(table.cell(7, 1))


# Add text to each cells
table.cell(0, 0).text = "Job Number " + str(jobNumber)
table.cell(1, 0).text = "Client:"
table.cell(3, 0).text = "Password:"
table.cell(4, 0).text = "Address:"
table.cell(5, 0).text = "Issue:"
table.cell(6, 0).text = "Items Serviced:"
table.cell(7, 0).text = "To Invoice"

table.cell(0, 1).text = "Date:"
table.cell(1, 1).text = "Phone:"
table.cell(2, 1).text = "Email:"
table.cell(3, 1).text = "Parts:"
table.cell(5, 1).text = "Work Done:" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "Data Saved?  Y / N"

# Add a new column, split in the middle
section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')

# Add a space between two tables
paragraph = document.add_paragraph()

jobNumber +=1

# Add a new table
table = document.add_table(rows=8, cols=2, style='Table Grid')

# Change row height
table.rows[0].height = Cm(1.1)
table.rows[1].height = Cm(1.1)
table.rows[2].height = Cm(1.1)
table.rows[3].height = Cm(1.1)
table.rows[4].height = Cm(1.1)
table.rows[5].height = Cm(1.1)
table.rows[6].height = Cm(1.1)
table.rows[7].height = Cm(1.1)

table.columns[0].width = Cm(6.8)
table.columns[1].width = Cm(6.8)

table.cell(1, 0).merge(table.cell(2, 0))
table.cell(4, 0).merge(table.cell(4, 1))
table.cell(5, 1).merge(table.cell(7, 1))


# Add text to each cells
table.cell(0, 0).text = "Job Number " + str(jobNumber)
table.cell(1, 0).text = "Client:"
table.cell(3, 0).text = "Password:"
table.cell(4, 0).text = "Address:"
table.cell(5, 0).text = "Issue:"
table.cell(6, 0).text = "Items Serviced:"
table.cell(7, 0).text = "To Invoice"

table.cell(0, 1).text = "Date:"
table.cell(1, 1).text = "Phone:"
table.cell(2, 1).text = "Email:"
table.cell(3, 1).text = "Parts:"
table.cell(5, 1).text = "Work Done:" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "Data Saved?  Y / N"

paragraph = document.add_paragraph()

jobNumber +=1

# Add a new table
table = document.add_table(rows=8, cols=2, style='Table Grid')

# Change row height
table.rows[0].height = Cm(1.1)
table.rows[1].height = Cm(1.1)
table.rows[2].height = Cm(1.1)
table.rows[3].height = Cm(1.1)
table.rows[4].height = Cm(1.1)
table.rows[5].height = Cm(1.1)
table.rows[6].height = Cm(1.1)
table.rows[7].height = Cm(1.1)

table.columns[0].width = Cm(6.8)
table.columns[1].width = Cm(6.8)

table.cell(1, 0).merge(table.cell(2, 0))
table.cell(4, 0).merge(table.cell(4, 1))
table.cell(5, 1).merge(table.cell(7, 1))


# Add text to each cells
table.cell(0, 0).text = "Job Number " + str(jobNumber)
table.cell(1, 0).text = "Client:"
table.cell(3, 0).text = "Password:"
table.cell(4, 0).text = "Address:"
table.cell(5, 0).text = "Issue:"
table.cell(6, 0).text = "Items Serviced:"
table.cell(7, 0).text = "To Invoice"

table.cell(0, 1).text = "Date:"
table.cell(1, 1).text = "Phone:"
table.cell(2, 1).text = "Email:"
table.cell(3, 1).text = "Parts:"
table.cell(5, 1).text = "Work Done:" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "Data Saved?  Y / N"

paragraph = document.add_paragraph()

jobNumber +=1

# Add a new table
table = document.add_table(rows=8, cols=2, style='Table Grid')

# Change row height
table.rows[0].height = Cm(1.1)
table.rows[1].height = Cm(1.1)
table.rows[2].height = Cm(1.1)
table.rows[3].height = Cm(1.1)
table.rows[4].height = Cm(1.1)
table.rows[5].height = Cm(1.1)
table.rows[6].height = Cm(1.1)
table.rows[7].height = Cm(1.1)

table.columns[0].width = Cm(6.8)
table.columns[1].width = Cm(6.8)

table.cell(1, 0).merge(table.cell(2, 0))
table.cell(4, 0).merge(table.cell(4, 1))
table.cell(5, 1).merge(table.cell(7, 1))


# Add text to each cells
table.cell(0, 0).text = "Job Number " + str(jobNumber)
table.cell(1, 0).text = "Client:"
table.cell(3, 0).text = "Password:"
table.cell(4, 0).text = "Address:"
table.cell(5, 0).text = "Issue:"
table.cell(6, 0).text = "Items Serviced:"
table.cell(7, 0).text = "To Invoice"

table.cell(0, 1).text = "Date:"
table.cell(1, 1).text = "Phone:"
table.cell(2, 1).text = "Email:"
table.cell(3, 1).text = "Parts:"
table.cell(5, 1).text = "Work Done:" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "\n" \
                        "Data Saved?  Y / N"

# save jobNumber in file.
file = open('var.txt', 'wb')
pickle.dump(jobNumber, file)
file.close()

document.save('JobCard.docx')


path = ("JobCard.docx")